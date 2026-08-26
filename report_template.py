"""
TEN Capital — Pitch Deck Analysis Report Template
=================================================

The document skeleton every generated analysis follows. This module owns the
*structure and formatting only*; all content is supplied by the caller through
the `ReportData` mapping described below. No company-specific text lives here.

Usage
-----
    from report_template import build_report, blank_report_data

    data = blank_report_data()          # fully-formed skeleton with placeholders
    data["company_name"] = "..."        # populate from the model output
    ...
    build_report(Path("out.docx"), data)

Document map
------------
    Header block    Company / Document title / Source / Date
    SECTION 1 — PITCH DECK ANALYSIS: STRENGTHS, WEAKNESSES & RECOMMENDATIONS
        1.1  Section-by-Section Assessment  Section | Strengths | Weaknesses | Recs
        1.2  Formatting, Storytelling & Design Recommendations
        1.3  Proposed Revised Slide Outline Slide | Content
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- Brand palette -------------------------------------------------------------
# Matches the TEN Capital web app's palette (navy / coral / amber / teal), tuned
# for legibility on a white page rather than the app's dark UI.

NAVY = "121B2E"        # section headings, table header fill, primary brand
TEAL = "128C86"        # sub-headings, document subtitle, accent text
GREY = "5C6E86"        # metadata lines, footer
ZEBRA_FILL = "EAF6F5"  # alternating body-row fill (teal tint)
WHITE = "FFFFFF"

# Three-colour accent strip echoing the app's coral -> amber -> teal gradient.
ACCENT_STRIP = ("EE5A4E", "F3A22A", "35BEBB")

# --- Typography --------------------------------------------------------------

BODY_FONT = "Arial"
FOOTER_FONT = "Open Sans"

SIZE_TITLE = Pt(22)
SIZE_SUBTITLE = Pt(13)
SIZE_META = Pt(10)
SIZE_SECTION = Pt(15)
SIZE_SUBSECTION = Pt(11)
SIZE_BODY = Pt(10)
SIZE_TABLE = Pt(9)
SIZE_FOOTER = Pt(7)

PAGE_MARGIN = Inches(0.75)
FOOTER_LOGO = "TEN_Capital_logo_footer.png"  # resolved relative to this file

# --- The framework being applied ---------------------------------------------
# Deck sections assessed in 1.1 — fixed rows of the strengths/weaknesses table.

DECK_SECTIONS: list[str] = [
    "Executive Summary / Cover",
    "Problem Statement",
    "Solution / Technology",
    "Market Opportunity",
    "Business Model",
    "Team",
    "Traction & Milestones",
    "Competitive Landscape",
    "Risk & Mitigation",
    "Deal Terms / Investment Ask",
    "Call-to-Action / Closing",
]

# Design dimensions covered in 1.2 — one labelled paragraph each.
DESIGN_DIMENSIONS: list[str] = [
    "TYPOGRAPHY",
    "NARRATIVE FLOW",
    "DATA VISUALIZATION",
    "SLIDE DENSITY",
    "COLOR CONSISTENCY",
]


# --- Low-level formatting helpers --------------------------------------------


def _shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def _write_cell(
    cell,
    text: str,
    *,
    fill: str = WHITE,
    color: str = "000000",
    bold: bool = False,
    size: Pt = SIZE_TABLE,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    _shade(cell, fill)
    para = cell.paragraphs[0]
    para.text = ""
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(str(text))
    run.font.name = BODY_FONT
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _add_table(doc: Document, headers: Sequence[str], widths: Sequence[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_borders(table)
    for i, header in enumerate(headers):
        _write_cell(table.rows[0].cells[i], header, fill=NAVY, color=WHITE, bold=True)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def _add_body_row(table, values: Sequence[str], index: int):
    """Append a zebra-striped body row."""
    fill = WHITE if index % 2 == 0 else ZEBRA_FILL
    cells = table.add_row().cells
    for i, value in enumerate(values):
        _write_cell(cells[i], value, fill=fill)
    return cells


def _heading(doc: Document, text: str, *, size: Pt, color: str, before: int, after: int):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = size
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    return para


def section_heading(doc: Document, text: str):
    """SECTION N — TITLE"""
    return _heading(doc, text, size=SIZE_SECTION, color=NAVY, before=18, after=9)


def subsection_heading(doc: Document, text: str):
    """N.N  Sub-section title"""
    return _heading(doc, text, size=SIZE_SUBSECTION, color=TEAL, before=12, after=6)


def body_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(5)
    if bold_prefix:
        run = para.add_run(f"{bold_prefix} ")
        run.font.name = BODY_FONT
        run.font.size = SIZE_BODY
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    run = para.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = SIZE_BODY
    return para


# --- Fixed document furniture ------------------------------------------------


def _add_header_accent(doc: Document) -> None:
    """Thin coral -> amber -> teal strip at the top of every page, echoing the app UI."""
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.header_distance = Inches(0.25)
        table = section.header.add_table(rows=1, cols=len(ACCENT_STRIP), width=Inches(1))
        table.autofit = False
        no_borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            no_borders.append(element)
        table._tbl.tblPr.append(no_borders)

        usable_width = section.page_width - section.left_margin - section.right_margin
        each = usable_width // len(ACCENT_STRIP)
        row = table.rows[0]
        height = row._tr.get_or_add_trPr()
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), "120")  # 6pt, in twentieths of a point
        tr_height.set(qn("w:hRule"), "exact")
        height.append(tr_height)

        for i, color in enumerate(ACCENT_STRIP):
            cell = row.cells[i]
            cell.width = each
            _shade(cell, color)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(1)


def _add_footer(doc: Document, document_title: str, compiled_on: str) -> None:
    """Single-line centred TEN Capital footer: title · page · compiled-on · logo."""
    logo = _find_logo()
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        para = section.footer.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _run(text: str):
            run = para.add_run(text)
            run.font.name = FOOTER_FONT
            run.font.size = SIZE_FOOTER
            run.font.color.rgb = RGBColor.from_string(GREY)
            return run

        _run(f"{document_title}          ")
        _add_page_number(_run(""))
        _run(f"     Compiled on {compiled_on} by TEN Capital Network    ")
        if logo is not None:
            para.add_run().add_picture(str(logo), width=Inches(0.67), height=Inches(0.25))


def _find_logo() -> Path | None:
    """Locate the footer logo beside this module, or in the parent folder."""
    here = Path(__file__).parent
    for candidate in (here / FOOTER_LOGO, here.parent / FOOTER_LOGO):
        if candidate.exists():
            return candidate
    return None


def _add_page_number(run) -> None:
    begin, field, end = OxmlElement("w:fldChar"), OxmlElement("w:instrText"), OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field.set(qn("xml:space"), "preserve")
    field.text = "PAGE"
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, field, end):
        run._r.append(element)


def header_block(doc: Document, data: dict) -> None:
    """Compact centred header — company, document title, source, date — sitting
    directly atop Section 1 rather than on its own mostly-blank cover page."""

    def centred(text: str, size: Pt, color: str, bold: bool, after: int):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(after)
        run = para.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = size
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)

    centred(data["company_name"].upper(), SIZE_TITLE, NAVY, True, 8)
    centred(data["document_title"], SIZE_SUBTITLE, TEAL, False, 6)
    centred(f"Source: {data['source']}", SIZE_META, GREY, False, 2)
    centred(f"{data['date']}  |  Prepared by TEN Capital Network", SIZE_META, GREY, False, 14)


# --- Section builders ---------------------------------------------------------


def section_1_deck_analysis(doc: Document, data: dict) -> None:
    section_heading(doc, "SECTION 1 — PITCH DECK ANALYSIS: STRENGTHS, WEAKNESSES & RECOMMENDATIONS")
    body_paragraph(
        doc,
        "The following analysis applies the deck review framework: a section-by-section "
        "assessment of what is working, what is hurting investor perception, and specific "
        "actionable improvements.",
    )

    subsection_heading(doc, "1.1  Section-by-Section Assessment")
    table = _add_table(doc, ["Section", "Strengths", "Weaknesses / Gaps", "Recommendations"],
                       widths=[1.3, 1.9, 1.9, 1.9])
    for i, row in enumerate(data["section_assessment"]):
        cells = _add_body_row(
            table,
            [row["section"], row["strengths"], row["weaknesses"], row["recommendations"]],
            i,
        )
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
    doc.add_paragraph()

    subsection_heading(doc, "1.2  Formatting, Storytelling & Design Recommendations")
    for item in data["design_recommendations"]:
        body_paragraph(doc, item["text"], bold_prefix=f"{item['label']}:")

    subsection_heading(doc, "1.3  Proposed Revised Slide Outline")
    table = _add_table(doc, ["Slide", "Content"], widths=[0.8, 6.2])
    for i, row in enumerate(data["revised_outline"]):
        _add_body_row(table, [row["slide"], row["content"]], i)


# --- Entry point --------------------------------------------------------------


def build_report(output_path: Path | str, data: dict) -> Path:
    """Render `data` into the branded TEN Capital pitch deck analysis document."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = SIZE_BODY
    for section in doc.sections:
        section.left_margin = section.right_margin = PAGE_MARGIN
        section.top_margin = section.bottom_margin = PAGE_MARGIN

    _add_header_accent(doc)

    header_block(doc, data)
    section_1_deck_analysis(doc, data)

    _add_footer(doc, data["document_title"], data["date"])

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


# --- Skeleton data ------------------------------------------------------------

PLACEHOLDER = "[TO BE COMPLETED]"


def blank_report_data() -> dict:
    """
    A fully-formed, content-free instance of the data contract.

    Every key the template reads is present, pre-expanded to the fixed rows of
    the framework (11 deck sections, 5 design dimensions) so callers can fill
    values in place rather than reconstructing the shape.
    """
    return {
        # Header block
        "company_name": PLACEHOLDER,
        "document_title": "TEN Capital Pitch Deck Analysis",
        "source": PLACEHOLDER,                      # source deck filename
        "date": PLACEHOLDER,                        # e.g. "January 1, 2026"

        # 1.1 / 1.2 / 1.3
        "section_assessment": [
            {
                "section": name,
                "strengths": PLACEHOLDER,
                "weaknesses": PLACEHOLDER,
                "recommendations": PLACEHOLDER,
            }
            for name in DECK_SECTIONS
        ],
        "design_recommendations": [
            {"label": label, "text": PLACEHOLDER} for label in DESIGN_DIMENSIONS
        ],
        "revised_outline": [{"slide": n, "content": PLACEHOLDER} for n in range(1, 17)],
    }


def validate_report_data(data: dict) -> list[str]:
    """Return a list of structural problems; empty list means the data fits the template."""
    problems: list[str] = []

    def require(container: dict, keys: Iterable[str], where: str) -> None:
        for key in keys:
            if key not in container:
                problems.append(f"missing key '{key}' in {where}")

    require(data, blank_report_data().keys(), "report data")
    if not problems:
        for name, expected in (
            ("section_assessment", len(DECK_SECTIONS)),
            ("design_recommendations", len(DESIGN_DIMENSIONS)),
        ):
            if len(data[name]) != expected:
                problems.append(f"'{name}' has {len(data[name])} rows, expected {expected}")
        if not data["revised_outline"]:
            problems.append("'revised_outline' is empty")
    return problems


if __name__ == "__main__":
    # Emit the empty skeleton so the layout can be inspected without an API call.
    out = build_report(Path(__file__).with_name("_template_preview.docx"), blank_report_data())
    print(f"Wrote template preview -> {out}")
